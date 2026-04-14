#!/usr/bin/env python3
"""
Phase 27: brutal product smoke checks (API-path behavior, not the eval harness).

This script exercises the same workflow a real user uses:
  upload (simulated by writing files) -> sync -> simple chat -> delete doc -> chat again

It runs **without starting FastAPI** by calling the shared domain services directly.
Use this to catch mismatches between "benchmark eval passes" and "live workflow fails".
"""

from __future__ import annotations

import os
import time
from pathlib import Path


def _write_long_txt(path: Path) -> None:
    blocks = []
    blocks.append("# Q4 Operations Review\n")
    blocks.append(
        "This is a long multi-section report with headings, lists, and some messy line wraps.\n"
        "It intentionally includes anchors for testing: ZEPHYR-PROD, p99, latency, and a finance table.\n"
    )
    blocks.append("\n## Reliability and performance\n")
    blocks.append(
        "QueryPlane expectations: p99 latency should be below 400 ms for typical single-user workloads.\n"
        "If p99 latency exceeds budget, operators should reduce top_k or chunk size.\n"
    )
    blocks.append("\n## Finance snapshot\n")
    blocks.append("Line-wrap stress:\nRevenue for Q3 was\n47,000,000 USD\nand CEO is Jane Okonkwo.\n")
    blocks.append("\nTABLE\nItem | Value\nRevenue | 47,000,000\nCEO | Jane Okonkwo\n")
    # Bulk filler to emulate a longer report
    for i in range(60):
        blocks.append(f"\n### Appendix {i}\nNotes: ZEPHYR-PROD section {i}.")
    path.write_text("\n".join(blocks), encoding="utf-8")


def _write_docx(path: Path) -> None:
    from docx import Document

    d = Document()
    d.add_heading("Internal Playbook (DOCX)", level=1)
    d.add_paragraph("This DOCX exists to test extraction order and tables.")
    d.add_heading("Latency requirements", level=2)
    d.add_paragraph("p99 latency should remain below 400 milliseconds.")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Target"
    table.cell(1, 0).text = "p99 latency"
    table.cell(1, 1).text = "400ms"
    d.save(path)


def main() -> int:
    from app.services import index_service
    from app.services.chat_service import answer_user_query
    from app.services.library_delete import delete_library_document

    base = Path("eval/_phase27_smoke_work").resolve()
    raw = base / "raw"
    faiss = base / "faiss"
    raw.mkdir(parents=True, exist_ok=True)
    faiss.mkdir(parents=True, exist_ok=True)

    # Make sure we are testing API-like behavior: do NOT auto-sync on chat.
    os.environ.pop("KA_AUTO_SYNC_ON_CHAT", None)
    os.environ.setdefault("KA_NO_STREAM", "1")

    txt = raw / "long_report.txt"
    docx = raw / "table_doc.docx"
    _write_long_txt(txt)
    _write_docx(docx)

    t0 = time.perf_counter()
    ok, msg, nvec, action = index_service.rebuild_knowledge_index(raw, faiss, chunk_size=900, chunk_overlap=120)
    dt = time.perf_counter() - t0
    print(f"[sync] ok={ok} action={action} vectors={nvec} took_s={dt:.2f} msg={msg!r}")
    if not ok:
        return 2

    # Broad overview
    turn = answer_user_query(
        "What is this document about, in plain language?",
        raw_dir=raw,
        faiss_folder=faiss,
        chunk_size=900,
        chunk_overlap=120,
        top_k=6,
    )
    print(f"[chat] broad mode={turn.mode} hits={len(turn.hits or [])} preview={turn.text[:120]!r}")

    # Narrow factual
    turn2 = answer_user_query(
        "What does it say about p99 latency requirements?",
        raw_dir=raw,
        faiss_folder=faiss,
        chunk_size=900,
        chunk_overlap=120,
        top_k=6,
    )
    print(f"[chat] latency mode={turn2.mode} hits={len(turn2.hits or [])} preview={turn2.text[:120]!r}")

    # Delete one doc and ensure chat prompts for sync rather than silently using stale index.
    okd, msgd = delete_library_document(raw, faiss, "long_report.txt", chunk_size=900, chunk_overlap=120)
    print(f"[delete] ok={okd} msg={msgd!r}")

    turn3 = answer_user_query(
        "What is this document about?",
        raw_dir=raw,
        faiss_folder=faiss,
        chunk_size=900,
        chunk_overlap=120,
        top_k=6,
    )
    print(f"[chat] after_delete mode={turn3.mode} note={turn3.assistant_note!r} hits={len(turn3.hits or [])}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

