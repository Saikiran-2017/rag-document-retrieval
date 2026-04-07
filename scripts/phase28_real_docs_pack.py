#!/usr/bin/env python3
"""
Phase 28: real-document workflow pack (separate from benchmark ``eval/gold_cases.json``).

Builds a messy/long TXT, a DOCX with headings and tables, and (if reportlab is installed)
a multi-page text PDF. Then rebuilds a temp FAISS index and runs representative queries via
``answer_user_query``. Does not start FastAPI.

Install PDF support for this script::

    pip install reportlab

Usage (repository root)::

    set PYTHONPATH=.
    python scripts/phase28_real_docs_pack.py

With OpenAI (full chat checks)::

    set OPENAI_API_KEY=...
    python scripts/phase28_real_docs_pack.py
"""

from __future__ import annotations

import json
import os
import sys
import time
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


def _write_messy_long_txt(path: Path) -> None:
    """Simulates messy reflow: hyphen breaks, section headers, wrapped numbers."""
    lines = [
        "# Northwind Reliability Playbook",
        "",
        "Executive summary: this document describes ZEPHYR-PROD uptime goals and finance figures.",
        "",
        "## Latency SLO",
        "The platform target for p99 latency is below four hundred milliseconds for interactive",
        "queries. Operators must alert if p99 excee-",
        "ds the budget for more than fifteen minutes.",
        "",
        "## Finance (messy wrap)",
        "Annual run rate was",
        "52,500,000",
        "USD under CFO Maria Chen.",
        "",
        "### Deep section anchor",
        "PHASE28-ANCHOR: section seven discusses disaster recovery.",
    ]
    for i in range(80):
        lines.append(f"\n### Appendix block {i}\nFiller line with token RETAIN-{i} for length.")
    path.write_text("\n".join(lines), encoding="utf-8")


def _write_structured_docx(path: Path) -> None:
    from docx import Document

    d = Document()
    d.add_heading("Operations manual (DOCX)", level=1)
    d.add_paragraph("Intro paragraph for Phase 28 DOCX extraction.")
    d.add_heading("Tables and metrics", level=2)
    d.add_paragraph("p99 latency must stay under 400 ms when measured at the edge.")
    table = d.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Target"
    table.cell(1, 0).text = "p99 latency"
    table.cell(1, 1).text = "400ms"
    table.cell(2, 0).text = "Owner"
    table.cell(2, 1).text = "SRE on-call"
    d.save(path)


def _try_write_pdf_reportlab(path: Path) -> bool:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        return False
    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    y = height - 72
    for i, line in enumerate(
        [
            "Phase 28 PDF — selectable text",
            "Keyword: ZEPHYR-PROD reliability charter.",
            "p99 latency budget: 400 milliseconds.",
            "CFO name on record: Maria Chen.",
        ]
    ):
        c.drawString(72, y - i * 16, line)
    c.showPage()
    c.drawString(72, height - 72, "Page two — PHASE28-PDF-PAGE2 anchor for section queries.")
    c.save()
    return True


def _run_queries(raw: Path, faiss: Path) -> list[dict[str, object]]:
    from app.services.chat_service import answer_user_query

    qs = [
        ("broad_overview", "What is this document about in plain language?"),
        ("summary", "Summarize the main operational themes across my files."),
        ("section", "What does section seven or disaster recovery say?"),
        ("narrow_fact", "What is the p99 latency target in milliseconds?"),
        ("cfo_fact", "Who is named as CFO or finance lead?"),
    ]
    out: list[dict[str, object]] = []
    for qid, qtext in qs:
        t0 = time.perf_counter()
        turn = answer_user_query(
            qtext,
            raw_dir=raw,
            faiss_folder=faiss,
            chunk_size=900,
            chunk_overlap=120,
            top_k=6,
        )
        dt = time.perf_counter() - t0
        hits = len(turn.hits or [])
        text = (turn.text or "")[:400]
        out.append(
            {
                "id": qid,
                "mode": turn.mode,
                "hits": hits,
                "seconds": round(dt, 3),
                "text_preview": text,
                "assistant_note": (turn.assistant_note or "")[:200] or None,
            }
        )
    return out


def main() -> int:
    os.environ.pop("KA_AUTO_SYNC_ON_CHAT", None)
    os.environ.setdefault("KA_NO_STREAM", "1")

    base = Path("eval/_phase28_real_docs_work").resolve()
    raw = base / "raw"
    faiss = base / "faiss"
    raw.mkdir(parents=True, exist_ok=True)
    faiss.mkdir(parents=True, exist_ok=True)

    _write_messy_long_txt(raw / "messy_long_report.txt")
    _write_structured_docx(raw / "structured_tables.docx")
    pdf_ok = _try_write_pdf_reportlab(raw / "phase28_text.pdf")
    if not pdf_ok:
        print("[phase28] reportlab not installed — skipping PDF fixture (pip install reportlab).")

    from app.ingestion.loader import load_file
    from app.services import index_service

    # Extraction sanity (no API key)
    for name in ["messy_long_report.txt", "structured_tables.docx"]:
        docs = load_file(raw / name)
        assert docs, f"expected text from {name}"
        total = sum(len(d.text) for d in docs)
        print(f"[extract] {name} segments={len(docs)} chars={total}")

    if pdf_ok:
        pd = load_file(raw / "phase28_text.pdf")
        print(f"[extract] phase28_text.pdf segments={len(pd)} chars={sum(len(d.text) for d in pd)}")

    t0 = time.perf_counter()
    ok, msg, nvec, action = index_service.rebuild_knowledge_index(
        raw, faiss, chunk_size=900, chunk_overlap=120
    )
    print(f"[sync] ok={ok} action={action} vectors={nvec} took_s={time.perf_counter() - t0:.2f} msg={msg!r}")
    if not ok:
        return 2

    key = os.environ.get("OPENAI_API_KEY", "").strip()
    results: list[dict[str, object]] = []
    if not key:
        print("[phase28] OPENAI_API_KEY unset — skipping LLM query pack (index built OK).")
    else:
        results = _run_queries(raw, faiss)
        for row in results:
            print(f"[chat] {row['id']} mode={row['mode']} hits={row['hits']} t={row['seconds']}s")

    report = {
        "pack": "phase28_real_documents",
        "pdf_fixture": pdf_ok,
        "index_ok": ok,
        "vector_count": nvec,
        "queries": results,
    }
    out_path = Path("eval/_report_phase28_real_docs.json")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(f"[phase28] wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
