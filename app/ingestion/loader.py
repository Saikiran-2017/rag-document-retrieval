"""
Load local documents into structured `IngestedDocument` records.

Each record carries clean text plus metadata for downstream chunking and retrieval.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Literal

import pdfplumber
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from app.ingestion.pdf_ocr_optional import try_ocr_pdf_page

logger = logging.getLogger(__name__)

FileType = Literal["pdf", "docx", "txt"]

SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".docx", ".txt"})


@dataclass(frozen=True)
class IngestedDocument:
    """
    One logical piece of content after ingestion (e.g. one PDF page or one whole TXT file).

    Downstream steps can split `text` into smaller chunks while preserving `metadata_dict`.
    """

    text: str
    source_name: str
    absolute_path: Path
    file_type: FileType
    page_number: int | None
    document_id: str
    metadata_dict: dict[str, Any] = field(init=False, repr=False)

    def __post_init__(self) -> None:
        object.__setattr__(
            self,
            "metadata_dict",
            {
                "source_name": self.source_name,
                "file_path": str(self.absolute_path.resolve()),
                "file_type": self.file_type,
                "page_number": self.page_number,
                "document_id": self.document_id,
            },
        )


def get_default_raw_dir() -> Path:
    """Project `data/raw` directory (next to `app/`)."""
    return Path(__file__).resolve().parent.parent.parent / "data" / "raw"


def _normalize_line(raw: str | None) -> str:
    """Collapse horizontal whitespace on one visual line."""
    if not raw:
        return ""
    return " ".join(raw.split())


def _normalize_text(raw: str | None) -> str:
    """
    Normalize whitespace while preserving paragraph and line structure.

    Blank lines separate paragraphs (joined with ``\\n\\n``). Non-blank lines that
    were separated only by a single newline stay as ``\\n`` within a paragraph
    (helps tables and reflowed PDF lines).
    """
    if not raw:
        return ""
    t = raw.replace("\r\n", "\n").replace("\r", "\n")
    norm_lines = [_normalize_line(ln) for ln in t.split("\n")]
    blocks: list[str] = []
    cur: list[str] = []
    for w in norm_lines:
        if w == "":
            if cur:
                blocks.append("\n".join(cur))
                cur = []
        else:
            cur.append(w)
    if cur:
        blocks.append("\n".join(cur))
    return "\n\n".join(blocks)


def _format_docx_table(table: DocxTable) -> str:
    """Flatten a DOCX table to newline-separated rows, cells as 'a | b'."""
    rows_out: list[str] = []
    try:
        for row in table.rows:
            cells = [_normalize_line(c.text) for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                rows_out.append(" | ".join(cells))
    except Exception as exc:
        logger.debug("DOCX table extraction issue: %s", exc)
        return ""
    return "\n".join(rows_out)


def _pypdf_page_text(path: Path, page_index_0: int) -> str:
    """Second-chance text for one PDF page (0-based index)."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        reader = PdfReader(str(path))
        if page_index_0 < 0 or page_index_0 >= len(reader.pages):
            return ""
        return reader.pages[page_index_0].extract_text() or ""
    except Exception as exc:
        logger.debug("pypdf page %s of %s: %s", page_index_0 + 1, path.name, exc)
        return ""


def _file_type_for_path(path: Path) -> FileType:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix == ".txt":
        return "txt"
    raise ValueError(
        f"Unsupported file type '{suffix}' for {path.name}. "
        f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def _load_pdf_pages(path: Path) -> list[IngestedDocument]:
    """Extract one record per non-empty page (1-based page numbers)."""
    source_name = path.name
    file_type: FileType = "pdf"
    stem = path.stem
    results: list[IngestedDocument] = []

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            raw = ""
            try:
                raw = page.extract_text() or ""
            except Exception as exc:
                logger.warning("PDF page %s in %s failed to extract: %s", i, path.name, exc)

            text = _normalize_text(raw)
            if not text:
                raw2 = _pypdf_page_text(path, i - 1)
                text = _normalize_text(raw2)
            if not text:
                raw3 = try_ocr_pdf_page(path, i)
                text = _normalize_text(raw3) if raw3 else ""

            if not text:
                continue

            doc_id = f"{stem}_pdf_p{i}"
            results.append(
                IngestedDocument(
                    text=text,
                    source_name=source_name,
                    absolute_path=path,
                    file_type=file_type,
                    page_number=i,
                    document_id=doc_id,
                )
            )

    return results


def _load_docx_whole(path: Path) -> list[IngestedDocument]:
    """Load DOCX in document order: paragraphs and tables interleaved."""
    source_name = path.name
    file_type: FileType = "docx"
    stem = path.stem

    document = DocxDocument(path)
    parts: list[str] = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para = DocxParagraph(child, document)
            t = _normalize_line(para.text)
            if t:
                parts.append(t)
        elif child.tag == qn("w:tbl"):
            tbl = DocxTable(child, document)
            block = _format_docx_table(tbl)
            if block:
                parts.append(block)

    text = "\n\n".join(parts)
    text = _normalize_text(text)
    if not text:
        return []

    doc_id = f"{stem}_docx"
    return [
        IngestedDocument(
            text=text,
            source_name=source_name,
            absolute_path=path,
            file_type=file_type,
            page_number=None,
            document_id=doc_id,
        )
    ]


def _load_txt_whole(path: Path) -> list[IngestedDocument]:
    """Load a text file as UTF-8 (invalid bytes replaced so Windows files do not crash)."""
    source_name = path.name
    file_type: FileType = "txt"
    stem = path.stem

    raw = path.read_text(encoding="utf-8", errors="replace")
    text = _normalize_text(raw)
    if not text:
        return []

    doc_id = f"{stem}_txt"
    return [
        IngestedDocument(
            text=text,
            source_name=source_name,
            absolute_path=path,
            file_type=file_type,
            page_number=None,
            document_id=doc_id,
        )
    ]


def load_file(path: str | Path) -> list[IngestedDocument]:
    """
    Load a single file and return zero or more `IngestedDocument` rows.

    PDFs yield one row per non-empty page. DOCX and TXT yield at most one row
    (none if the document has no extractable text).
    """
    p = Path(path).expanduser().resolve()
    if not p.is_file():
        raise FileNotFoundError(f"Not a file: {p}")

    ft = _file_type_for_path(p)
    if ft == "pdf":
        return _load_pdf_pages(p)
    if ft == "docx":
        return _load_docx_whole(p)
    return _load_txt_whole(p)


def load_raw_directory(raw_dir: str | Path | None = None) -> list[IngestedDocument]:
    """
    Load every supported file in a directory (non-recursive).

    Unknown extensions are skipped. A file that errors is logged and skipped so
    one bad document does not stop the whole batch.
    """
    base = Path(raw_dir).expanduser().resolve() if raw_dir is not None else get_default_raw_dir()
    if not base.is_dir():
        raise NotADirectoryError(f"Not a directory: {base}")

    all_docs: list[IngestedDocument] = []
    for path in sorted(base.iterdir()):
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            all_docs.extend(load_file(path))
        except Exception as exc:
            logger.warning("Skipping %s: %s", path.name, exc)

    return all_docs


def print_ingestion_summary(documents: list[IngestedDocument]) -> None:
    """Print a human-readable summary for manual testing."""
    print(f"Ingested segments: {len(documents)}")
    for i, doc in enumerate(documents, start=1):
        preview = doc.text[:120] + ("…" if len(doc.text) > 120 else "")
        print(f"\n--- [{i}] {doc.document_id} ---")
        print(f"  source: {doc.source_name}")
        print(f"  type:   {doc.file_type}  page: {doc.page_number}")
        print(f"  chars:  {len(doc.text)}")
        print(f"  preview: {preview!r}")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    raw = get_default_raw_dir()
    print(f"Loading from: {raw}\n")
    loaded = load_raw_directory(raw)
    print_ingestion_summary(loaded)
    if not loaded:
        print(
            "\n(No documents found. Add .pdf, .docx, or .txt files under data/raw/ and run again.)"
        )
