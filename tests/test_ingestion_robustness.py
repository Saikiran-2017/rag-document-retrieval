"""Phase 22: normalization, chunking, PDF fallback, and difficult fixtures (no API keys)."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

ROOT = Path(__file__).resolve().parents[1]
DIFFICULT_FIXTURE = ROOT / "fixtures" / "difficult_mixed_layout.txt"


def test_normalize_preserves_paragraphs_and_intraline_breaks():
    from app.ingestion import loader

    raw = "Line one\nLine two\n\nSecond block\n\n"
    out = loader._normalize_text(raw)
    assert "Line one\nLine two" in out
    assert "\n\n" in out
    assert "Second block" in out


def test_difficult_fixture_loads_and_chunks_for_retrieval_anchors():
    from app.ingestion.loader import load_file
    from app.utils.chunker import chunk_ingested_documents

    assert DIFFICULT_FIXTURE.is_file()
    docs = load_file(DIFFICULT_FIXTURE)
    assert len(docs) == 1
    full = docs[0].text
    assert "ZEPHYR-TEST-ANCHOR-PH22" in full
    assert "Widget A" in full and "12.4M" in full
    assert "## Markdown style section" in full

    chunks = chunk_ingested_documents(docs, chunk_size=220, chunk_overlap=40)
    joined = "\n".join(c.text for c in chunks)
    assert "ZEPHYR-TEST-ANCHOR-PH22" in joined
    assert any("latency" in c.text for c in chunks)


def test_format_docx_table_mock():
    from app.ingestion.loader import _format_docx_table

    def cell(t: str) -> MagicMock:
        m = MagicMock()
        m.text = t
        return m

    row1 = MagicMock()
    row1.cells = [cell("Revenue"), cell("Q3")]
    row2 = MagicMock()
    row2.cells = [cell("Total"), cell("99")]
    tbl = MagicMock()
    tbl.rows = [row1, row2]
    assert _format_docx_table(tbl) == "Revenue | Q3\nTotal | 99"


def test_docx_roundtrip_with_table(tmp_path: Path):
    from docx import Document

    from app.ingestion.loader import load_file

    p = tmp_path / "t.docx"
    d = Document()
    d.add_paragraph("Intro PH22DOCX")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "KPI"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "p99"
    table.cell(1, 1).text = "400ms"
    d.add_paragraph("After table")
    d.save(p)

    rows = load_file(p)
    assert len(rows) == 1
    t = rows[0].text
    assert "Intro PH22DOCX" in t
    assert "KPI" in t and "400ms" in t
    assert "After table" in t
    assert t.find("Intro") < t.find("KPI") < t.find("After")


def test_pdf_pypdf_fallback_when_pdfplumber_page_empty(tmp_path: Path):
    from app.ingestion.loader import load_file

    p = tmp_path / "one.pdf"
    p.write_bytes(b"%PDF-1.4\n")

    page = MagicMock()
    page.extract_text.return_value = ""
    inner = MagicMock()
    inner.pages = [page]
    ctx = MagicMock()
    ctx.__enter__.return_value = inner
    ctx.__exit__.return_value = None

    with (
        patch("app.ingestion.loader.pdfplumber.open", return_value=ctx),
        patch(
            "app.ingestion.loader._pypdf_page_text",
            return_value="Recovered PDF text for page one.",
        ),
        patch("app.ingestion.loader.try_ocr_pdf_page", return_value=""),
    ):
        out = load_file(p)

    assert len(out) == 1
    assert "Recovered PDF" in out[0].text


def test_pdf_ocr_disabled_by_default(monkeypatch: pytest.MonkeyPatch):
    monkeypatch.delenv("RAG_ENABLE_PDF_OCR", raising=False)
    from app.ingestion.pdf_ocr_optional import pdf_ocr_enabled, try_ocr_pdf_page

    assert pdf_ocr_enabled() is False
    assert try_ocr_pdf_page(Path("/nonexistent.pdf"), 1) == ""


def test_playbook_fixture_still_has_eval_keywords_after_normalize():
    """Regression guard: eval substring anchors survive ingestion."""
    from app.ingestion.loader import load_file

    pb = ROOT / "eval" / "fixtures" / "playbook_long.txt"
    docs = load_file(pb)
    assert len(docs) == 1
    t = docs[0].text
    for needle in ("ZEPHYR-7", "IngestPlane", "QueryPlane", "400", "p99"):
        assert needle in t
