"""Phase 28: PDF/DOCX normalization helpers."""

from __future__ import annotations

from pathlib import Path

from app.ingestion.loader import _repair_hyphen_line_breaks


def test_repair_hyphen_line_breaks() -> None:
    assert _repair_hyphen_line_breaks("foo-\nbar") == "foobar"
    assert _repair_hyphen_line_breaks("ok\nline") == "ok\nline"


def test_docx_heading_prefixes_in_output(tmp_path: Path) -> None:
    from docx import Document

    from app.ingestion.loader import load_file

    p = tmp_path / "h.docx"
    d = Document()
    d.add_heading("Title", level=1)
    d.add_paragraph("Body text.")
    d.add_heading("Section", level=2)
    d.add_paragraph("More.")
    d.save(p)
    docs = load_file(p)
    assert len(docs) == 1
    assert "# Title" in docs[0].text
    assert "## Section" in docs[0].text
